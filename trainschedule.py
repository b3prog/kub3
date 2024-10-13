import psycopg2
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime
import re
from docx import Document
from utils import center_window, RoundedButton




class TrainScheduleApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Расписание поездов")
        center_window(self.root, 0.8, 0.6)
        self.create_widgets()
        self.create_table()
        self.load_schedule()

    def connect_db(self):
        return psycopg2.connect(
            dbname="train_schedule",
            user="b3",
            password="b3b3",
            host="localhost"
        )

    def create_table(self):
        conn = self.connect_db()
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS trains (
                            id SERIAL PRIMARY KEY,
                            train_number VARCHAR(50),
                            departure_point VARCHAR(100),
                            destination VARCHAR(100),
                            departure_time TIMESTAMP,
                            arrival_time TIMESTAMP,
                            travel_time INTERVAL,
                            available_seats INTEGER)''')
        conn.commit()
        cursor.close()
        conn.close()

    def save_train(self, train_number, departure_point, destination, departure_time, arrival_time, travel_time, available_seats):
        conn = self.connect_db()
        cursor = conn.cursor()
        cursor.execute('''INSERT INTO trains (train_number, departure_point, destination, departure_time, arrival_time, travel_time, available_seats)
                          VALUES (%s, %s, %s, %s, %s, %s, %s)''',
                       (train_number, departure_point, destination, departure_time, arrival_time, travel_time, available_seats))
        conn.commit()
        cursor.close()
        conn.close()

    def load_schedule(self):
        conn = self.connect_db()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM trains')
        records = cursor.fetchall()
        cursor.close()
        conn.close()

        for row in self.tree.get_children():
            self.tree.delete(row)
        for record in records:
            departure_time = record[4].strftime("%H:%M %d-%m-%Y")
            arrival_time = record[5].strftime("%H:%M %d-%m-%Y")
            travel_time = str(record[6])

            self.tree.insert('', 'end', values=(
                record[1], record[2], record[3], departure_time, arrival_time, travel_time, record[7]))

    def delete_train(self, train_number):
        conn = self.connect_db()
        cursor = conn.cursor()
        cursor.execute('DELETE FROM trains WHERE train_number = %s', (train_number,))
        conn.commit()
        cursor.close()
        conn.close()

    def update_seats(self, train_number, available_seats):
        conn = self.connect_db()
        cursor = conn.cursor()
        cursor.execute('UPDATE trains SET available_seats = %s WHERE train_number = %s', (available_seats, train_number))
        conn.commit()
        cursor.close()
        conn.close()
    def create_word_file(self):
        # Создаем новый документ
        doc = Document()
        doc.add_heading('Расписание поездов', level=5)

        # Добавляем таблицу в документ
        columns = ['Номер поезда', 'Пункт отправления', 'Пункт назначения', 'Время отправления',
                   'Время прибытия', 'Время в пути', 'Свободные места']

        # Создаем таблицу в документе Word (строка для заголовков + данные)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'

        # Заполняем заголовки таблицы
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(columns):
            hdr_cells[i].text = column_name

        # Добавляем данные из Treeview в таблицу
        for row in self.tree.get_children():
            values = self.tree.item(row)['values']
            row_cells = table.add_row().cells
            for i, value in enumerate(values):
                row_cells[i].text = str(value)

        # Добавляем дату в название файла
        current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        file_name = f"train_schedule_{current_time}.docx"

        # Сохраняем документ
        doc.save(file_name)
        messagebox.showinfo("Готово", f"Файл успешно создан: {file_name}")

    def create_widgets(self):
        style = ttk.Style()
        style.configure('Treeview', font=('bahnschrift semilight', 10), rowheight=25)
        style.configure('Treeview.Heading', font=('bahnschrift bold', 12, ))

        columns = (
        'номер_поезда', 'пункт_отправления', 'пункт_назначения', 'время_отправления', 'время_прибытия', 'время_в_пути',
        'свободные_места')

        self.tree = ttk.Treeview(self.root, columns=columns, show='headings')
        for col in columns:
            self.tree.heading(col, text=col.replace('_', ' ').title())
        self.tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        button_frame = tk.Frame(self.root)
        button_frame.pack(fill=tk.X)

        # Создаем закругленные кнопки
        RoundedButton(button_frame, text="Добавить поезд", command=self.add_train).pack(side=tk.LEFT, padx=5, pady=5,
                                                                                        expand=True)
        RoundedButton(button_frame, text="Удалить поезд", command=self.delete_train_button).pack(side=tk.LEFT, padx=5,
                                                                                                 pady=5, expand=True)
        RoundedButton(button_frame, text="Найти поезд", command=self.find_train_button).pack(side=tk.LEFT, padx=5,
                                                                                             pady=5, expand=True)
        RoundedButton(button_frame, text="Изменить места", command=self.update_seats_button).pack(side=tk.LEFT, padx=5,
                                                                                                  pady=5, expand=True)
        RoundedButton(button_frame, text="Сформировать файл", command=self.create_word_file).pack(side=tk.LEFT, padx=5,
                                                                                                  pady=5, expand=True)

    def format_time_input(self, entry, format_type):
        text = entry.get()
        if format_type == "departure_arrival":
            formatted = re.sub(r"[^0-9]", "", text)
            if len(formatted) > 2:
                formatted = formatted[:2] + ':' + formatted[2:]
            if len(formatted) > 5:
                formatted = formatted[:5] + ' ' + formatted[5:]
            if len(formatted) > 8:
                formatted = formatted[:8] + ':' + formatted[8:]
            if len(formatted) > 11:
                formatted = formatted[:11] + ':' + formatted[11:]
            entry.delete(0, tk.END)
            entry.insert(0, formatted[:16])

        elif format_type == "travel_time":
            formatted = re.sub(r"[^0-9]", "", text)
            if len(formatted) > 2:
                formatted = formatted[:2] + ':' + formatted[2:]
            if len(formatted) > 5:
                formatted = formatted[:5] + ':' + formatted[5:]
            entry.delete(0, tk.END)
            entry.insert(0, formatted[:8])

    def add_train(self):
        add_window = tk.Toplevel(self.root)
        center_window(add_window, 0.23, 0.235)  # Центрируем окно
        add_window.grab_set()  # Блокируем основное окно

        add_window.title("Добавить поезд")

        label_font = ("bahnschrift semilight", 12)
        # Метки с изменённым шрифтом
        tk.Label(add_window, text="№ поезда:", font=label_font).grid(row=0, column=0, padx=5, pady=2)
        tk.Label(add_window, text="Пункт отправления", font=label_font).grid(row=1, column=0, padx=5, pady=2)
        tk.Label(add_window, text="Пункт назначения", font=label_font).grid(row=2, column=0, padx=5, pady=2)
        tk.Label(add_window, text="Время отправления (чч:мм дд:мм:гггг)", font=label_font).grid(row=3, column=0, padx=5, pady=2)
        tk.Label(add_window, text="Время прибытия (чч:мм дд:мм:гггг)", font=label_font).grid(row=4, column=0, padx=5, pady=2)
        tk.Label(add_window, text="Время в пути (дд:чч:мм)", font=label_font).grid(row=5, column=0, padx=5, pady=2)
        tk.Label(add_window, text="Количество свободных мест", font=label_font).grid(row=6, column=0, padx=5, pady=2)

        train_number_entry = ttk.Entry(add_window, style="TEntry")
        departure_point_entry = ttk.Entry(add_window, style="TEntry")
        destination_entry = ttk.Entry(add_window, style="TEntry")
        departure_time_entry = ttk.Entry(add_window, style="TEntry")
        arrival_time_entry = ttk.Entry(add_window, style="TEntry")
        travel_time_entry = ttk.Entry(add_window, style="TEntry")
        available_seats_entry = ttk.Entry(add_window, style="TEntry")

        train_number_entry.grid(row=0, column=1, padx=5, pady=2)
        departure_point_entry.grid(row=1, column=1, padx=5, pady=2)
        destination_entry.grid(row=2, column=1, padx=5, pady=2)
        departure_time_entry.grid(row=3, column=1, padx=5, pady=2)
        arrival_time_entry.grid(row=4, column=1, padx=5, pady=2)
        travel_time_entry.grid(row=5, column=1, padx=5, pady=2)
        available_seats_entry.grid(row=6, column=1, padx=5, pady=2)

        # Bind formatting functions to relevant inputs
        departure_time_entry.bind("<KeyRelease>",
                                  lambda event: self.format_time_input(departure_time_entry, "departure_arrival"))
        arrival_time_entry.bind("<KeyRelease>",
                                lambda event: self.format_time_input(arrival_time_entry, "departure_arrival"))
        travel_time_entry.bind("<KeyRelease>", lambda event: self.format_time_input(travel_time_entry, "travel_time"))

        def validate_input():
            train_number = train_number_entry.get()
            departure_point = departure_point_entry.get()
            destination = destination_entry.get()
            departure_time = departure_time_entry.get()
            arrival_time = arrival_time_entry.get()
            travel_time = travel_time_entry.get()
            available_seats = available_seats_entry.get()

            # Regular expressions for input validation
            russian_letters_regex = r'^[А-Яа-яЁё\-]+$'
            time_format_regex = r'^\d{2}:\d{2} \d{2}:\d{2}:\d{4}$'
            travel_time_format_regex = r'^\d+:\d{1,2}:\d{1,2}$'
            numbers_only_regex = r'^\d+$'

            # Validation checks
            if not re.match(numbers_only_regex, train_number):
                messagebox.showwarning("Ошибка", "Номер поезда должен содержать только числа")
                return False
            if not re.match(russian_letters_regex, departure_point):
                messagebox.showwarning("Ошибка", "Пункт отправления должен содержать только русские буквы и -")
                return False
            if not re.match(russian_letters_regex, destination):
                messagebox.showwarning("Ошибка", "Пункт назначения должен содержать только русские буквы и -")
                return False
            if not re.match(time_format_regex, departure_time):
                messagebox.showwarning("Ошибка", "Время отправления должно быть в формате чч:мм дд:мм:год")
                return False
            if not re.match(time_format_regex, arrival_time):
                messagebox.showwarning("Ошибка", "Время прибытия должно быть в формате чч:мм дд:мм:год")
                return False
            if not re.match(travel_time_format_regex, travel_time):
                messagebox.showwarning("Ошибка", "Время в пути должно быть в формате дней:часов:минут")
                return False
            if not re.match(numbers_only_regex, available_seats):
                messagebox.showwarning("Ошибка", "Количество свободных мест должно содержать только числа")
                return False

            return True

        def save_train():
            if validate_input():
                train_number = train_number_entry.get()
                departure_point = departure_point_entry.get()
                destination = destination_entry.get()
                departure_time = datetime.strptime(departure_time_entry.get(), "%H:%M %d:%m:%Y")
                arrival_time = datetime.strptime(arrival_time_entry.get(), "%H:%M %d:%m:%Y")
                travel_time = travel_time_entry.get()
                available_seats = int(available_seats_entry.get())

                self.save_train(train_number, departure_point, destination, departure_time, arrival_time, travel_time,
                                available_seats)
                self.load_schedule()

                # Добавляем сообщение о том, что поезд был успешно добавлен
                messagebox.showinfo("Успех", f"Поезд №{train_number} успешно добавлен!")

                add_window.destroy()

        RoundedButton(add_window, text="Сохранить", command=save_train, width=130, height=32).grid(row=7, columnspan=2, padx=10, pady=10)

    def delete_train_button(self):
        delete_window = tk.Toplevel(self.root)
        center_window(delete_window, 0.16, 0.09)  # Центрируем окно
        delete_window.grab_set()  # Блокируем основное окно

        delete_window.title("Удалить поезд")
        label_font = ("bahnschrift semilight", 12)
        tk.Label(delete_window, text="Введите номер поезда:",font=label_font).pack(pady=1)
        train_number_entry = ttk.Entry(delete_window)
        train_number_entry.pack()

        def confirm_delete():
            train_number = train_number_entry.get()
            if train_number:
                # Проверка, существует ли поезд
                conn = self.connect_db()
                cursor = conn.cursor()
                cursor.execute('SELECT * FROM trains WHERE train_number = %s', (train_number,))
                train = cursor.fetchone()
                cursor.close()
                conn.close()

                if train:
                    self.delete_train(train_number)
                    messagebox.showinfo("Успех", "Поезд удален!")
                    self.load_schedule()
                    delete_window.destroy()
                else:
                    messagebox.showerror("Ошибка", "Поезд с таким номером не найден.")
            else:
                messagebox.showwarning("Ошибка", "Введите номер поезда.")

        RoundedButton(delete_window, text="OK", command=confirm_delete, width=130, height=32).pack(side=tk.LEFT, padx=10, pady=1)
        RoundedButton(delete_window, text="Cancel", command=delete_window.destroy, width=130, height=32).pack(side=tk.RIGHT, padx=10, pady=1)

    def find_train_button(self):
        find_window = tk.Toplevel(self.root)
        center_window(find_window, 0.16, 0.09)  # Центрируем окно
        find_window.grab_set()  # Блокируем основное окно

        find_window.title("Найти поезд")
        label_font = ("bahnschrift semilight", 12)
        tk.Label(find_window, text="Введите пункт назначения:",font=label_font).pack(pady=1)
        destination_entry = ttk.Entry(find_window)
        destination_entry.pack()

        def find_train():
            destination = destination_entry.get()
            if destination:
                trains = []
                for train in self.tree.get_children():
                    item = self.tree.item(train)
                    if item['values'][2] == destination:
                        trains.append(item['values'])
                if trains:
                    trains_info = "\n".join(
                        [f"№ поезда: {train[0]}, Время отправления: {train[3]}" for train in trains])
                    messagebox.showinfo("Найденные поезда", trains_info)
                else:
                    messagebox.showinfo("Результат", "Поезда не найдены.")
                find_window.destroy()

        RoundedButton(find_window, text="OK", command=find_train, width=130, height=32).pack(side=tk.LEFT, padx=10, pady=1)
        RoundedButton(find_window, text="Cancel", command=find_window.destroy, width=130, height=32).pack(side=tk.RIGHT, padx=10, pady=1)

    def update_seats_button(self):
        update_window = tk.Toplevel(self.root)
        center_window(update_window, 0.16, 0.13)  # Центрируем окно
        update_window.grab_set()  # Блокируем основное окно

        update_window.title("Изменить места")
        label_font = ("bahnschrift semilight", 12)
        tk.Label(update_window, text="Введите номер поезда:",font=label_font).pack()
        train_number_entry = ttk.Entry(update_window)
        train_number_entry.pack()

        tk.Label(update_window, text="Новое количество свободных мест:",font=label_font).pack()
        available_seats_entry = ttk.Entry(update_window)
        available_seats_entry.pack()

        def update_seats():
            train_number = train_number_entry.get()
            available_seats = available_seats_entry.get()
            if train_number and available_seats.isdigit():
                # Проверка на существование поезда
                conn = self.connect_db()
                cursor = conn.cursor()
                cursor.execute('SELECT * FROM trains WHERE train_number = %s', (train_number,))
                train = cursor.fetchone()
                cursor.close()
                conn.close()

                if train:
                    self.update_seats(train_number, int(available_seats))
                    messagebox.showinfo("Успех", "Количество свободных мест обновлено!")
                    self.load_schedule()
                    update_window.destroy()
                else:
                    messagebox.showerror("Ошибка", "Поезд с таким номером не найден.")
            else:
                messagebox.showerror("Ошибка", "Введите корректные данные.")

        RoundedButton(update_window, text="OK", command=update_seats, width=130, height=32).pack(side=tk.LEFT, padx=10, pady=1)
        RoundedButton(update_window, text="Cancel", command=update_window.destroy, width=130, height=32).pack(side=tk.RIGHT, padx=10, pady=1)

